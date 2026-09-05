"""
Build FINAL_INTRODUCTION_TO_RAG_LESSON.docx from the approved pipeline output.

Reads output/lesson_output.md — the accepted lesson from the passing run —
and renders it as a formatted Word document. If the lesson file is a
diagnostic draft (the run did not clear the quality bar), the script refuses
to overwrite the submission document.

Run from the project root:

    python assets/create_docx.py

Optional infographic placement: the two PNGs in assets/ are inserted after
the "How does RAG work" and "Concrete Example" sections when present.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
LESSON_PATH = PROJECT_ROOT / "output" / "lesson_output.md"
OUTPUT_PATH = PROJECT_ROOT / "FINAL_INTRODUCTION_TO_RAG_LESSON.docx"
INFographic_1 = PROJECT_ROOT / "assets" / "infographic1_rag_flow.png"
INFographic_2 = PROJECT_ROOT / "assets" / "infographic2_comparison.png"


def parse_markdown(text: str):
    """Yield (kind, payload) blocks from the lesson Markdown."""
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            yield ("heading", (level, stripped[level:].strip()))
        elif re.match(r"^[-*]\s+", stripped):
            yield ("bullet", _inline(stripped))
        elif re.match(r"^\d+[.)]\s+", stripped):
            yield ("numbered", _inline(stripped))
        else:
            yield ("para", _inline(stripped))
        i += 1


def _inline(md: str):
    """Split 'text' into (bold_prefix, rest) when a bullet starts with **Term**."""
    m = re.match(r"^[-*]\s+\*\*(.+?)\*\*\s*:?\s*(.*)$", md)
    if m:
        return ("term", m.group(1), m.group(2))
    return ("text", None, md)


def _clean(md_text: str) -> str:
    """Strip markdown emphasis for plain runs."""
    t = re.sub(r"\*\*(.+?)\*\*", r"\1", md_text)
    t = re.sub(r"\*(.+?)\*", r"\1", t)
    return t.replace("`", "")


def _style_run(run, size=11):
    run.font.name = "Arial"
    run.font.size = Pt(size)


def add_infographic(doc, image_path: Path, width_in: float):
    if not image_path.is_file():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width_in))


def build_docx(lesson_text: str) -> None:
    doc = Document()

    title = doc.add_heading("Introduction to RAG", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = "Arial"

    pending_infographic = None  # insert after the matching section heading

    for kind, payload in parse_markdown(lesson_text):
        if kind == "heading":
            level, text = payload
            doc_level = min(level, 3)
            h = doc.add_heading(_clean(text), level=doc_level)
            for run in h.runs:
                run.font.name = "Arial"
            low = text.lower()
            if "how" in low and "work" in low:
                pending_infographic = ("flow", 3.5)
            elif "example" in low:
                pending_infographic = ("compare", 5.0)
        elif kind == "bullet":
            style, term, rest = payload
            p = doc.add_paragraph(style="List Bullet")
            if style == "term" and rest:
                r1 = p.add_run(f"{term}: ")
                r1.bold = True
                _style_run(r1)
                r2 = p.add_run(_clean(rest))
                _style_run(r2)
            else:
                r = p.add_run(_clean(rest if style == "term" else payload[2]))
                _style_run(r)
        elif kind == "numbered":
            _, term, rest = payload
            p = doc.add_paragraph(style="List Number")
            r = p.add_run(_clean(rest if term is None else f"{term} {rest}"))
            _style_run(r)
        else:  # para
            _, term, rest = payload
            p = doc.add_paragraph()
            r = p.add_run(_clean(rest))
            _style_run(r)

        if pending_infographic and kind != "heading":
            name, width = pending_infographic
            add_infographic(doc, INFographic_1 if name == "flow" else INFographic_2, width)
            pending_infographic = None

    doc.save(str(OUTPUT_PATH))


def main() -> None:
    if not LESSON_PATH.is_file():
        print(f"ERROR: {LESSON_PATH} not found. Run the pipeline first.")
        sys.exit(1)

    lesson_text = LESSON_PATH.read_text(encoding="utf-8")
    if "DIAGNOSTIC DRAFT" in lesson_text:
        print(
            "ERROR: output/lesson_output.md is a diagnostic draft (the last run "
            "did not clear the quality bar). The submission document is only "
            "built from an ACCEPTED lesson. Re-run the pipeline until it passes."
        )
        sys.exit(1)

    build_docx(lesson_text)
    print(f"DOCX written to {OUTPUT_PATH} from the accepted lesson.")


if __name__ == "__main__":
    main()
